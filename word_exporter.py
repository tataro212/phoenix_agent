"""
Word Document Exporter for Phoenix Agent
Converts translated blueprint to Microsoft Word (.docx) format.
"""

import sys
from pathlib import Path
from typing import Dict, Any, Optional
from document_blueprint import load_blueprint, DocumentBlueprint, DocumentElement

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. Install with: pip install python-docx")


class WordExporter:
    """Exports translated blueprint to Microsoft Word format."""
    
    def __init__(self, blueprint: DocumentBlueprint):
        self.blueprint = blueprint
        self.doc = Document()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """Setup document styles for consistent formatting."""
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        h1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        h2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Heading 3 style
        h3_style = self.doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Arial'
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Normal paragraph style
        normal_style = self.doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
    
    def add_heading_with_bookmark(self, text: str, level: int, bookmark_id: str):
        """Add a heading with a bookmark for navigation."""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.style = self.doc.styles['CustomHeading1']
        elif level == 2:
            heading = self.doc.add_heading(text, level=2)
            heading.style = self.doc.styles['CustomHeading2']
        elif level == 3:
            heading = self.doc.add_heading(text, level=3)
            heading.style = self.doc.styles['CustomHeading3']
        else:
            heading = self.doc.add_heading(text, level=level)
        
        # Add bookmark (if supported)
        try:
            run = heading.runs[0] if heading.runs else heading.add_run()
            self.add_bookmark(run, bookmark_id)
        except Exception as e:
            print(f"⚠️  Could not add bookmark for {bookmark_id}: {e}")
    
    def add_bookmark(self, run, bookmark_id: str):
        """Add a bookmark to a run."""
        tag = run._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_id)
        tag.append(start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        tag.append(end)
    
    def export_to_word(self, output_path: str) -> bool:
        """Export the blueprint to a Word document."""
        if not DOCX_AVAILABLE:
            print("❌ Error: python-docx not available. Install with: pip install python-docx")
            return False
        
        try:
            # Add document title
            if self.blueprint.get('metadata', {}).get('title'):
                title = self.doc.add_heading(self.blueprint['metadata']['title'], level=0)
                title.style = self.doc.styles['CustomTitle']
            else:
                title = self.doc.add_heading("Translated Document", level=0)
                title.style = self.doc.styles['CustomTitle']
            
            # Process each page
            for page_data in self.blueprint['pages']:
                page_num = page_data['page_number']
                print(f"Processing page {page_num} for Word export...")
                
                # Get reading order for this page
                reading_order = page_data['reading_order']
                elements = {el['id']: el for el in page_data['elements']}
                
                # Process elements in reading order
                for element_id in reading_order:
                    if element_id in elements:
                        element = elements[element_id]
                        self.process_element(element)
                
                # Add page break if not the last page
                if page_num < len(self.blueprint['pages']):
                    self.doc.add_page_break()
            
            # Save the document
            self.doc.save(output_path)
            print(f"✅ Successfully exported document to: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error exporting to Word: {e}")
            return False
    
    def process_element(self, element: DocumentElement):
        """Process a single element for Word export."""
        element_type = element['type']
        content = element['content']
        
        if element_type == 'heading_1':
            self.add_heading_with_bookmark(content, 1, element['id'])
        
        elif element_type == 'heading_2':
            self.add_heading_with_bookmark(content, 2, element['id'])
        
        elif element_type == 'heading_3':
            self.add_heading_with_bookmark(content, 3, element['id'])
        
        elif element_type == 'paragraph':
            paragraph = self.doc.add_paragraph(content)
            paragraph.style = self.doc.styles['CustomNormal']
            
            # Apply formatting based on properties
            properties = element.get('properties', {})
            if properties.get('font_weight') == 'bold':
                for run in paragraph.runs:
                    run.bold = True
            if properties.get('is_italic'):
                for run in paragraph.runs:
                    run.italic = True
        
        elif element_type == 'list_item':
            # Handle list items with proper indentation
            properties = element.get('properties', {})
            list_level = properties.get('list_level', 1)
            
            # Create bullet point
            if properties.get('list_style') == 'bullet':
                paragraph = self.doc.add_paragraph(content, style='List Bullet')
            elif properties.get('list_style') == 'numbered':
                paragraph = self.doc.add_paragraph(content, style='List Number')
            else:
                paragraph = self.doc.add_paragraph(content)
            
            # Apply indentation for nested lists
            if list_level > 1:
                paragraph.paragraph_format.left_indent = Inches(0.25 * list_level)
        
        elif element_type == 'table':
            # Add table placeholder
            paragraph = self.doc.add_paragraph(f"[TABLE: {content}]")
            paragraph.style = self.doc.styles['CustomNormal']
        
        elif element_type == 'image':
            # Add image placeholder
            paragraph = self.doc.add_paragraph(f"[IMAGE: {content}]")
            paragraph.style = self.doc.styles['CustomNormal']
        
        elif element_type == 'caption':
            # Add caption with smaller font
            paragraph = self.doc.add_paragraph(content)
            paragraph.style = self.doc.styles['CustomNormal']
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.italic = True
        
        elif element_type in ['header', 'footer']:
            # Add header/footer content
            paragraph = self.doc.add_paragraph(content)
            paragraph.style = self.doc.styles['CustomNormal']
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = None  # Gray color


def export_blueprint_to_word(blueprint_path: str, output_path: str) -> bool:
    """
    Main function to export a blueprint to Word format.
    
    Args:
        blueprint_path: Path to the translated blueprint JSON file
        output_path: Path where the Word document should be saved
    
    Returns:
        bool: True if successful, False otherwise
    """
    # Load blueprint
    blueprint = load_blueprint(blueprint_path)
    if not blueprint:
        print(f"❌ Failed to load blueprint: {blueprint_path}")
        return False
    
    # Create exporter and export
    exporter = WordExporter(blueprint)
    return exporter.export_to_word(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python word_exporter.py <translated_blueprint.json> [output.docx]")
        sys.exit(1)
    
    blueprint_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "translated_document.docx"
    
    success = export_blueprint_to_word(blueprint_path, output_path)
    if not success:
        sys.exit(1) 